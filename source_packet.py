from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

def highlight_text(paragraph, text_to_highlight):
    for run in paragraph.runs:
        if text_to_highlight in run.text:
            start_index = run.text.find(text_to_highlight)
            end_index = start_index + len(text_to_highlight)
            new_run = paragraph.add_run(run.text[start_index:end_index])
            new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

def create_highlighted_document(original_source, derivative):
    doc = Document()

    derivative_paragraphs = derivative.split("\n")

    for derivative_paragraph in derivative_paragraphs:
        doc.add_paragraph(derivative_paragraph)

    for paragraph in doc.paragraphs:
        for derivative_paragraph in derivative_paragraphs:
            if derivative_paragraph in paragraph.text:
                highlight_text(paragraph, derivative_paragraph)

    return doc

def save_document(document, filename):
    document.save(filename)

if __name__ == "__main__":
    original_source = """
    President Joe Biden on Thursday told Israeli Prime Minister Benjamin Netanyahu that the overall humanitarian situation in Gaza is unacceptable and warned Israel to take steps to address the crisis or face consequences, a stark statement from Israel’s staunchest ally.

The 30-minute conversation was the two leaders’ first phone call since an Israeli strike killed seven aid workers from the World Central Kitchen in Gaza. The killings have set off a furor inside the White House and Biden has been described as reaching a new level of frustration with Israel’s campaign.

“President Biden emphasized that the strikes on humanitarian workers and the overall humanitarian situation are unacceptable,” the White House said in a statement shortly after the call wrapped. “He made clear that US policy with respect to Gaza will be determined by our assessment of Israel’s immediate action on these steps.”

Biden also said Israel needed to “announce and implement a series of specific, concrete, and measurable steps to address civilian harm, humanitarian suffering, and the safety of aid workers.”

Speaking in Brussels after the call, Secretary of State Antony Blinken made the stakes clear: “If we don’t see the changes that we need to see, there’ll be changes in our own policy,” he said. However, neither Blinken nor National Security Council spokesman John Kirby, speaking later at the White House press briefing, detailed what those potential policy changes could be.

Kirby said the US “would hope to see some announcements of changes here in the coming hours and days – and I’ll leave it at that.”

The call amounts to perhaps the most serious sign of Biden’s frustration with Israel’s campaign in Gaza, which was launched in the wake of the October 7 attacks by Hamas. In the time since, more than 32,000 people have been killed in Gaza, according to the Palestinian enclave’s health ministry, and grisly reports of civilian deaths are made every day.

The war has become one of Biden’s chief domestic political problems ahead of November’s election as key parts of his voter coalition have been outraged by the president’s support for Israel’s war. Protests have popped up at nearly every public event the president has held outside the White House in recent months and he was confronted in an intimate meeting with Muslim leaders earlier this week with stark opposition to the US’ policy toward the war.

That political conundrum made Thursday’s call and the statements from Blinken and Kirby particularly notable, as both men indicated possible change in the US’ posture toward the war may come if Israel keeps up its current practices.

Blinken said Israel should not stoop to the level of Hamas in its response to the group’s October 7 attacks.

“Israel is not Hamas,” Blinken said. “Israel is a democracy, Hamas a terrorist organization. And democracies place the highest value on human life - every human life.”

Biden was shaken by the strike on the World Central Kitchen workers, Kirby said.

Israel has acknowledged responsibility for the strikes but said the convoy was not targeted and the workers’ deaths were not intentional. The country continues to investigate the circumstances surrounding the killings. Blinken acknowledged the World Central Kitchen strike was not the first time Israel has killed aid workers in the conflict.

“It must be the last,” he said.

But even at the same time that the administration was alluding to potential changes in policy if Israel doesn’t stop killing civilians, the US was sending more deadly weapons to its ally.

Biden is set to green light an $18 billion sale of fighter jets from the United States and Israel and the administration recently authorized the transfer to Israel of over 1,000 500-pound bombs and over 1,000 small-diameter bombs, according to three people familiar with the matter.

Kirby defended the arms sale and transfers as the product of yearslong processes.

“With the exception of the immediate two months after the attack, we haven’t really sent emergency aid and military assistance to Israel,” Kirby said. “What you’re seeing here is the result of a process of foreign military sales to Israel that takes years.”

Earlier on Thursday, one of Biden’s closest allies on Capitol Hill – Democratic Sen. Chris Coons – told CNN that for the first time he is open to applying conditions to aid for Israel after the strike on the World Central Kitchen workers.

“I think we’re at that point,” he told CNN’s Sara Sidner. “I think we’re at the point where President Biden has said – and I have said and others have said – if Benjamin Netanyahu, prime minister, were to order the IDF into Rafah at scale, they were to drop thousand pound bombs and send in a battalion to go after Hamas and make no provision for civilians or for humanitarian aid, that I would vote to condition aid to Israel.”

The president has called for a temporary ceasefire that includes the release of hostages held by Hamas, and has repeatedly said he does not want Israel to undertake a ground invasion of Rafah in southern Gaza, where scores of displaced Gazans have sought shelter. He again called for a temporary ceasefire in the call with Netanyahu on Thursday.

“He underscored that an immediate ceasefire is essential to stabilize and improve the humanitarian situation and protect innocent civilians, and he urged the prime minister to empower his negotiators to conclude a deal without delay to bring the hostages home,” the White House said in its statement following the call.

But, Biden has so far stopped short of calling for a permanent ceasefire. His hesitance to do so is increasingly out of step with the actions and public statements of other world leaders, including US allies.

This is a breaking story and will be updated

CNN’s Jennifer Hansler and Donald Judd contributed to this report.
    """

    derivative = """
    The recent phone call between President Joe Biden and Israeli Prime Minister Benjamin Netanyahu marks a significant shift in the relationship dynamics between the two allies. Biden's assertion of the unacceptability of the strikes on humanitarian workers in Gaza reflects a heightened level of frustration within the Biden administration regarding Israel's military campaign.

The White House's statement following the call underscores the seriousness of the situation, indicating that US policy towards Gaza will be contingent upon Israel's immediate actions to address civilian harm and humanitarian suffering. This stance suggests that the US may be considering changes in its policy towards Israel if these concerns are not adequately addressed.

Secretary of State Antony Blinken's remarks further emphasize the gravity of the situation, indicating that failure to see the necessary changes may result in alterations to US policy. Blinken's assertion that Israel, as a democracy, should uphold the highest value of human life underscores the moral dimension of the issue.

However, despite the rhetoric and expressions of concern, the US's continued arms sales to Israel raise questions about the consistency of its position. The sale of fighter jets and bombs to Israel, despite the ongoing humanitarian crisis, highlights the complexity of US-Israel relations and the challenges of balancing strategic interests with humanitarian concerns.

Senator Chris Coons' openness to applying conditions to aid for Israel reflects a growing sentiment within the Democratic Party to reassess the unconditional support traditionally extended to Israel. This shift in attitude suggests a potential divergence between the Biden administration and some members of Congress regarding the approach to US-Israel relations.

Overall, while President Biden's call for a temporary ceasefire and emphasis on humanitarian concerns demonstrates a willingness to engage diplomatically, the lack of a clear stance on a permanent ceasefire raises questions about the administration's long-term strategy and alignment with international consensus on resolving the conflict.

Please note that this analysis is based on the information available up to January 2022, and subsequent developments may have occurred since then.
    """

    highlighted_document = create_highlighted_document(original_source, derivative)
    save_document(highlighted_document, "highlighted_news.docx")
