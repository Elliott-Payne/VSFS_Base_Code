import openai
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

# Set your OpenAI API key
openai.api_key = ''

def chat_with_gpt(prompt, chat_history):
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "user", "content": prompt},
            {"role": "assistant", "content": chat_history}
        ]
    )

    return response['choices'][0]['message']['content']

# print("Welcome to ChatGPT. You can start chatting by typing your messages.")
# print("Type 'exit' to end the conversation.")

chat_history = ""

source = ""
product = ""

try:
    with open('source.txt', 'rb') as file:
        file_contents = file.read()
        # Decode the binary content with a specific encoding
        decoded_contents = file_contents.decode('utf-8')  # Replace 'utf-8' with the appropriate encoding
    source = decoded_contents
except FileNotFoundError:
    print("File not found.")
except UnicodeDecodeError:
    print("Error: Unable to decode the file with the specified encoding.")

try:
    with open('product.txt', 'rb') as file:
        file_contents = file.read()
        # Decode the binary content with a specific encoding
        decoded_contents = file_contents.decode('utf-8')  # Replace 'utf-8' with the appropriate encoding
    product = decoded_contents
except FileNotFoundError:
    print("File not found.")
except UnicodeDecodeError:
    print("Error: Unable to decode the file with the specified encoding.")

user_input = "This is the news source: " + source

response = chat_with_gpt(user_input, chat_history)
chat_history += user_input + "\n"
chat_history += response + "\n"
# print("ChatGPT:", response)

user_input = "This is the analysis of the new source: " + product

response = chat_with_gpt(user_input, chat_history)
chat_history += user_input + "\n"
chat_history += response + "\n"
# print("ChatGPT:\n", response)

user_input = """I want you to list all of the quotes from the news source that was referenced in the analysis. 
                Do not change or remove anything to the quotes from the news source. Do not include the quotations around the quotes.
                Do not include duplicate quotes. Do not add on anything besides the quotes. 
                The list must be numbered. The quotes should come in the order it shows up in the text"""


response = chat_with_gpt(user_input, chat_history)
chat_history += user_input + "\n"
chat_history += response + "\n"
print("ChatGPT:\n", response)

quotes_list = response.split("\n")
print(quotes_list)
print(len(quotes_list))

for i in range(len(quotes_list)):
    quotes_list[i] = quotes_list[i][4:len(quotes_list[i])-3]

print(quotes_list)

def create_word_doc(text, keywords, title, author, date):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_heading(author + " " + date, 1)
    paragraph = doc.add_paragraph("")
    text_index = 0
    for keyword in keywords:
        if keyword in text:
            index = text.find(keyword)
            paragraph.add_run(text[text_index: index])
            paragraph.add_run(text[index:index + len(keyword)]).font.highlight_color = WD_COLOR_INDEX.YELLOW 
            text_index = index + len(keyword)
    paragraph.add_run(text[text_index:])
    doc.save('highlighted_doc.docx')

create_word_doc(source, quotes_list, "Testing", "daniel tran", "12 15 2001")

