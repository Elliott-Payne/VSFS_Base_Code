from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

def highlight_text(run, text_to_highlight):
    if text_to_highlight in run.text:
        index = run.text.find(text_to_highlight)
        # start = run.text[:index]
        highlight_text = run.text[index :index + len(text_to_highlight)]
        # run.clear()
        run.add_text(text_to_highlight)
        run.add_text(end)
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW  # Use WD_COLOR_INDEX for highlight color

def create_word_doc(text, keywords):
    doc = Document()
    paragraph = doc.add_paragraph("")
    text_index = 0
    for keyword in keywords:
        if keyword in text:
            index = text.find(keyword)
            paragraph.add_run(text[text_index: index])
            paragraph.add_run(text[index:index + len(keyword)]).font.highlight_color = WD_COLOR_INDEX.YELLOW 
            text_index = index + len(keyword)
    paragraph.add_run(text[text_index:])
    doc.save('highlighted_doc.docx')

# Example usage:
text = """President Joe Biden on Thursday told Israeli Prime Minister Benjamin Netanyahu that the overall humanitarian situation in Gaza is unacceptable and warned Israel to take steps to address the crisis or face consequences, a stark statement from Israel’s staunchest ally.
The 30-minute conversation was the two leaders’ first phone call since an Israeli strike killed seven aid workers from the World Central Kitchen in Gaza. The killings have set off a furor inside the White House and Biden has been described as reaching a new level of frustration with Israel’s campaign.
“President Biden emphasized that the strikes on humanitarian workers and the overall humanitarian situation are unacceptable,” the White House said in a statement shortly after the call wrapped. “He made clear that US policy with respect to Gaza will be determined by our assessment of Israel’s immediate action on these steps.”
Biden also said Israel needed to “announce and implement a series of specific, concrete, and measurable steps to address civilian harm, humanitarian suffering, and the safety of aid workers.”
Speaking in Brussels after the call, Secretary of State Antony Blinken made the stakes clear: “If we don’t see the changes that we need to see, there’ll be changes in our own policy,” he said. However, neither Blinken nor National Security Council spokesman John Kirby, speaking later at the White House press briefing, detailed what those potential policy changes could be.
Kirby said the US “would hope to see some announcements of changes here in the coming hours and days – and I’ll leave it at that.”
The call amounts to perhaps the most serious sign of Biden’s frustration with Israel’s campaign in Gaza, which was launched in the wake of the October 7 attacks by Hamas. In the time since, more than 32,000 people have been killed in Gaza, according to the Palestinian enclave’s health ministry, and grisly reports of civilian deaths are made every day.
The war has become one of Biden’s chief domestic political problems ahead of November’s election as key parts of his voter coalition have been outraged by the president’s support for Israel’s war. Protests have popped up at nearly every public event the president has held outside the White House in recent months and he was confronted in an intimate meeting with Muslim leaders earlier this week with stark opposition to the US’ policy toward the war.
That political conundrum made Thursday’s call and the statements from Blinken and Kirby particularly notable, as both men indicated possible change in the US’ posture toward the war may come if Israel keeps up its current practices.
Blinken said Israel should not stoop to the level of Hamas in its response to the group’s October 7 attacks.
“Israel is not Hamas,” Blinken said. “Israel is a democracy, Hamas a terrorist organization. And democracies place the highest value on human life - every human life.”
Biden was shaken by the strike on the World Central Kitchen workers, Kirby said.
Israel has acknowledged responsibility for the strikes but said the convoy was not targeted and the workers’ deaths were not intentional. The country continues to investigate the circumstances surrounding the killings. Blinken acknowledged the World Central Kitchen strike was not the first time Israel has killed aid workers in the conflict.
“It must be the last,” he said.
But even at the same time that the administration was alluding to potential changes in policy if Israel doesn’t stop killing civilians, the US was sending more deadly weapons to its ally.
Biden is set to green light an $18 billion sale of fighter jets from the United States and Israel and the administration recently authorized the transfer to Israel of over 1,000 500-pound bombs and over 1,000 small-diameter bombs, according to three people familiar with the matter.
Kirby defended the arms sale and transfers as the product of yearslong processes.
“With the exception of the immediate two months after the attack, we haven’t really sent emergency aid and military assistance to Israel,” Kirby said. “What you’re seeing here is the result of a process of foreign military sales to Israel that takes years.”
Earlier on Thursday, one of Biden’s closest allies on Capitol Hill – Democratic Sen. Chris Coons – told CNN that for the first time he is open to applying conditions to aid for Israel after the strike on the World Central Kitchen workers.
“I think we’re at that point,” he told CNN’s Sara Sidner. “I think we’re at the point where President Biden has said – and I have said and others have said – if Benjamin Netanyahu, prime minister, were to order the IDF into Rafah at scale, they were to drop thousand pound bombs and send in a battalion to go after Hamas and make no provision for civilians or for humanitarian aid, that I would vote to condition aid to Israel.”
The president has called for a temporary ceasefire that includes the release of hostages held by Hamas, and has repeatedly said he does not want Israel to undertake a ground invasion of Rafah in southern Gaza, where scores of displaced Gazans have sought shelter. He again called for a temporary ceasefire in the call with Netanyahu on Thursday.
“He underscored that an immediate ceasefire is essential to stabilize and improve the humanitarian situation and protect innocent civilians, and he urged the prime minister to empower his negotiators to conclude a deal without delay to bring the hostages home,” the White House said in its statement following the call.
But, Biden has so far stopped short of calling for a permanent ceasefire. His hesitance to do so is increasingly out of step with the actions and public statements of other world leaders, including US allies.
"""

# keywords = ['emphasized that the strikes on humanitarian workers and the overall humanitarian situation are unacceptable,']
keywords = ['President Biden emphasized that the strikes on humanitarian workers and the overall humanitarian situation are unacceptable,',
            'He made clear that US policy with respect to Gaza will be determined by our assessment of Israel’s immediate action on these steps.', 
            'Biden also said Israel needed to “announce and implement a series of specific, concrete, and measurable steps to address civilian harm, humanitarian suffering, and the safety of aid workers.', 
            'If we don’t see the changes that we need to see, there’ll be changes in our own policy,', 
            'Israel is not Hamas," Blinken said. "Israel is a democracy, Hamas a terrorist organization. And democracies place the highest value on human life - every human life.', 
            'It must be the last,', 
            'I think we’re at the point where President Biden has said – and I have said and others have said – if Benjamin Netanyahu, prime minister, were to order the IDF into Rafah at scale, they were to drop thousand pound bombs and send in a battalion to go after Hamas and make no provision for civilians or for humanitarian aid, that I would vote to condition aid to Israel.', 
            'He underscored that an immediate ceasefire is essential to stabilize and improve the humanitarian situation and protect innocent civilians, and he urged the prime minister to empower his negotiators to conclude a deal without delay to bring the hostages home,']

create_word_doc(text, keywords)